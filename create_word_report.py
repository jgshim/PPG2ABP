"""
Generate a detailed MS Word report for the PPG2ABP project.
"""
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ── Helper functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_table_with_style(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PPG 파형을 이용한\n침습적 동맥혈압 예측 모델 개발')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Prediction of Invasive Arterial Blood Pressure\nfrom Photoplethysmography in Pediatric Patients')
run.font.size = Pt(14)
run.italic = True
run.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('연구 보고서\n2026년 3월')
run.font.size = Pt(12)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
add_heading_styled('목차', level=1)
toc_items = [
    '1. 서론 (Introduction)',
    '   1.1 연구 배경',
    '   1.2 연구 목적',
    '   1.3 관련 연구',
    '2. 데이터 (Data)',
    '   2.1 데이터 개요',
    '   2.2 환자 인구통계 정보',
    '   2.3 VitalRecorder 파일 구조',
    '   2.4 신호 채널 정보',
    '3. 방법론 (Methods)',
    '   3.1 전체 파이프라인',
    '   3.2 데이터 전처리 (Step 1)',
    '   3.3 특징 추출 + 머신러닝 (Step 2)',
    '   3.4 1D-CNN 딥러닝 모델 (Step 3)',
    '   3.5 교차 검증 전략',
    '   3.6 평가 지표',
    '4. 결과 (Results)',
    '   4.1 데이터 추출 결과',
    '   4.2 머신러닝 모델 결과',
    '   4.3 1D-CNN 모델 결과',
    '   4.4 모델 비교',
    '   4.5 특징 중요도 분석',
    '5. 고찰 (Discussion)',
    '   5.1 결과 해석',
    '   5.2 한계점',
    '   5.3 향후 연구 방향',
    '6. 결론 (Conclusion)',
    '7. 참고문헌 (References)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════
add_heading_styled('1. 서론 (Introduction)', level=1)

add_heading_styled('1.1 연구 배경', level=2)
doc.add_paragraph(
    '동맥혈압(Arterial Blood Pressure, ABP)은 환자의 혈역학적 상태를 평가하는 가장 중요한 '
    '생체 신호 중 하나이다. 특히 수술 중이나 중환자실에서의 지속적인 혈압 모니터링은 환자 안전에 '
    '필수적이다. 현재 가장 정확한 연속 혈압 측정 방법은 동맥 카테터를 삽입하는 침습적 동맥혈압 '
    '측정(Invasive Blood Pressure, IBP)이지만, 이는 동맥 천자에 따른 감염, 혈전, 출혈 등의 '
    '합병증 위험이 있다.'
)
doc.add_paragraph(
    '반면 광용적맥파(Photoplethysmography, PPG)는 맥박산소측정기(Pulse Oximeter)를 통해 '
    '비침습적으로 간편하게 측정할 수 있는 신호로, 말초 혈관의 용적 변화를 반영한다. PPG 파형에는 '
    '혈관 탄성, 말초 혈관 저항, 심박출량 등 혈역학적 정보가 포함되어 있어, 이를 활용하여 동맥혈압을 '
    '추정하려는 시도가 활발히 이루어지고 있다.'
)
doc.add_paragraph(
    '그러나 기존 연구의 대부분은 성인 환자를 대상으로 수행되었으며, 소아 환자를 대상으로 한 연구는 '
    '매우 부족한 실정이다. 소아 환자는 성인과 비교하여 혈관 탄성, 심박수, 혈압 범위 등이 크게 다르며, '
    '연령에 따른 변화가 급격하여 별도의 모델 개발이 필요하다.'
)

add_heading_styled('1.2 연구 목적', level=2)
doc.add_paragraph('본 연구의 목적은 다음과 같다:')
objectives = [
    '소아 환자의 PPG 파형으로부터 침습적 동맥혈압(SBP, DBP, MBP)을 예측하는 모델을 개발한다.',
    '환자의 인구통계 정보(나이, 성별)를 보조 입력으로 활용하여 예측 정확도를 향상시킨다.',
    '전통적 머신러닝(XGBoost, LightGBM)과 딥러닝(1D-CNN) 접근법을 비교 분석한다.',
    '환자 단위 교차 검증(Leave-One-Subject-Out, Group K-Fold)을 통해 일반화 성능을 엄밀하게 평가한다.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

add_heading_styled('1.3 관련 연구', level=2)
doc.add_paragraph(
    'PPG를 이용한 혈압 추정 연구는 크게 두 가지 접근법으로 분류할 수 있다. '
    '첫째, PPG 파형에서 수작업으로 특징(feature)을 추출한 후 전통적 머신러닝 모델(Random Forest, '
    'SVR, XGBoost 등)로 혈압 수치를 예측하는 방법이다. 이 방법은 해석 가능성이 높고 적은 데이터에서도 '
    '비교적 안정적으로 작동하지만, 특징 설계에 도메인 지식이 필요하다는 한계가 있다.'
)
doc.add_paragraph(
    '둘째, 딥러닝 기반 접근법으로 원시 PPG 파형을 직접 입력하여 end-to-end로 학습하는 방법이다. '
    '1D-CNN, LSTM, U-Net, Transformer 등 다양한 아키텍처가 제안되었으며, 특히 1D-CNN 기반 '
    '모델이 PPG→ABP 변환에서 우수한 성능을 보여주고 있다. 그러나 대부분의 연구가 MIMIC 데이터베이스 '
    '등 성인 데이터를 사용하고 있어, 소아 대상 연구는 매우 제한적이다.'
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 2. DATA
# ════════════════════════════════════════════════════════════════
add_heading_styled('2. 데이터 (Data)', level=1)

add_heading_styled('2.1 데이터 개요', level=2)
doc.add_paragraph(
    '본 연구에서는 총 46명의 소아 환자로부터 수집된 수술 중 생체 신호 데이터를 사용하였다. '
    '데이터는 VitalRecorder 소프트웨어를 통해 GE Bx50 환자 모니터에서 기록되었으며, '
    '.vital 형식(gzip 압축 바이너리)으로 저장되었다.'
)
doc.add_paragraph(
    '각 파일에는 PPG 파형(PLETH), 침습적 동맥혈압 파형(IBP1), 심전도(ECG), 심박수(HR), '
    '호기말 이산화탄소(CO2) 등 다수의 생체 신호 채널이 포함되어 있다. 또한 모니터에서 자동 계산된 '
    '수축기혈압(ART1_SBP), 이완기혈압(ART1_DBP), 평균혈압(ART1_MBP) 수치값도 기록되었다.'
)

add_heading_styled('2.2 환자 인구통계 정보', level=2)
doc.add_paragraph(
    '환자의 성별 및 연령 정보는 별도의 엑셀 파일(PPG_ABP_cases.xlsx)에 기록되어 있다. '
    '전체 46명 중 남아 29명(63.0%), 여아 17명(37.0%)이며, 연령 범위는 0.25세(3개월)에서 '
    '15세까지이다.'
)

# Demographics table
demo_headers = ['항목', '값']
demo_rows = [
    ['전체 환자 수', '46명'],
    ['성별 (남/여)', '29명 (63.0%) / 17명 (37.0%)'],
    ['연령 범위', '0.25 ~ 15세'],
    ['연령 평균 ± 표준편차', '6.5 ± 3.5세'],
    ['데이터 형식', '.vital (gzip 압축 바이너리)'],
    ['모니터 장비', 'GE Bx50'],
]
add_table_with_style(demo_headers, demo_rows)
doc.add_paragraph()

# Full patient table
add_heading_styled('표 1. 전체 환자 목록', level=3)
patient_data = [
    ['case01','F','1'], ['case02','M','12'], ['case03','M','12'], ['case04','F','10'],
    ['case05','M','8'], ['case06','M','3'], ['case07','M','15'], ['case08','M','2'],
    ['case09','M','15'], ['case10','M','4'], ['case11','M','5'], ['case12','M','5'],
    ['case13','F','0.25'], ['case14','F','7'], ['case15','M','2'], ['case16','F','3'],
    ['case17','M','9'], ['case18','M','9'], ['case19','M','3'], ['case20','M','12'],
    ['case21','M','3'], ['case22','M','5'], ['case23','F','7'], ['case24','M','9'],
    ['case25','M','5'], ['case26','M','2'], ['case27','M','9'], ['case28','M','1'],
    ['case29','M','2'], ['case30','F','7'], ['case31','M','5'], ['case32','M','7'],
    ['case33','M','7'], ['case34','F','7'], ['case35','M','8'], ['case36','F','4'],
    ['case37','F','7'], ['case38','F','4'], ['case39','M','12'], ['case40','F','10'],
    ['case41','F','10'], ['case42','M','5'], ['case43','F','7'], ['case44','M','7'],
    ['case45','F','13'], ['case46','F','5'],
]
add_table_with_style(['Case ID', 'Sex', 'Age (yr)'], patient_data)
doc.add_paragraph()

add_heading_styled('2.3 VitalRecorder 파일 구조', level=2)
doc.add_paragraph(
    'VitalRecorder는 수술 중 환자 모니터, 마취기, BIS 등 다양한 의료 기기로부터 실시간으로 '
    '생체 신호를 기록하는 오픈소스 소프트웨어이다. .vital 파일은 gzip으로 압축된 바이너리 형식으로, '
    'Python의 vitaldb 패키지를 통해 읽을 수 있다.'
)
doc.add_paragraph(
    '각 트랙(track)은 "장비명/채널명" 형식으로 명명된다 (예: Bx50/PLETH, Bx50/IBP1). '
    '파형 데이터(waveform)는 고해상도(일반적으로 100~500Hz)로 기록되고, 수치 데이터(numeric)는 '
    '저해상도(일반적으로 1~2Hz)로 기록된다.'
)

add_heading_styled('2.4 신호 채널 정보', level=2)
channel_headers = ['채널명', '설명', '유형', '존재 파일 수']
channel_rows = [
    ['PLETH', 'Photoplethysmography (PPG)', 'Waveform', '46/46 (100%)'],
    ['IBP1', 'Invasive Blood Pressure', 'Waveform', '46/46 (100%)'],
    ['ECG1', 'Electrocardiogram', 'Waveform', '46/46 (100%)'],
    ['HR', 'Heart Rate', 'Numeric', '46/46 (100%)'],
    ['ART1_SBP', 'Systolic BP (auto-calculated)', 'Numeric', '46/46 (100%)'],
    ['ART1_DBP', 'Diastolic BP (auto-calculated)', 'Numeric', '46/46 (100%)'],
    ['ART1_MBP', 'Mean BP (auto-calculated)', 'Numeric', '46/46 (100%)'],
    ['CO2', 'Carbon Dioxide', 'Waveform', '46/46 (100%)'],
    ['NIBP_SBP/DBP', 'Non-Invasive BP', 'Numeric', '44/46 (95.7%)'],
    ['CVP2', 'Central Venous Pressure', 'Numeric', '3/46 (6.5%)'],
]
add_table_with_style(channel_headers, channel_rows)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 3. METHODS
# ════════════════════════════════════════════════════════════════
add_heading_styled('3. 방법론 (Methods)', level=1)

add_heading_styled('3.1 전체 파이프라인', level=2)
doc.add_paragraph(
    '본 연구의 분석 파이프라인은 3단계로 구성된다:'
)
steps = [
    'Step 1 (데이터 전처리): .vital 파일에서 PPG 및 ABP 파형을 추출하고, 필터링, 세그먼트 분할, '
    '품질 검사를 수행한다.',
    'Step 2 (특징 추출 + 머신러닝): PPG 세그먼트에서 37개의 시간/주파수 도메인 특징을 추출하고, '
    'XGBoost 및 LightGBM 모델로 SBP, DBP, MBP를 예측한다.',
    'Step 3 (1D-CNN 딥러닝): 원시 PPG 파형을 직접 입력으로 사용하는 1D Convolutional Neural '
    'Network를 학습시키고, 인구통계 정보(나이, 성별)를 Late Fusion으로 결합한다.',
]
for i, step in enumerate(steps):
    doc.add_paragraph(step, style='List Number')

add_heading_styled('3.2 데이터 전처리 (Step 1)', level=2)

doc.add_paragraph('3.2.1 리샘플링', style='Heading 4')
doc.add_paragraph(
    '원본 데이터의 샘플링 레이트는 장비에 따라 다를 수 있으므로, 모든 파형을 125Hz로 통일하여 '
    '리샘플링하였다. 125Hz는 PPG 분석에 널리 사용되는 표준 샘플링 레이트이다.'
)

doc.add_paragraph('3.2.2 세그먼트 분할', style='Heading 4')
doc.add_paragraph(
    '연속 파형을 10초 길이의 고정 윈도우로 분할하였다 (125Hz × 10초 = 1,250 샘플/세그먼트). '
    '각 세그먼트에 대해 다음 품질 기준을 적용하여 유효한 세그먼트만 선택하였다:'
)
quality_criteria = [
    '비결측(non-NaN) 비율 ≥ 95%: NaN이 5% 이상인 구간은 제외',
    'PPG 표준편차 > 0.1: 평탄 신호(flat signal) 제외',
    'IBP1 범위: 0~300 mmHg — 생리학적으로 불가능한 값 제외',
]
for c in quality_criteria:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph('3.2.3 필터링', style='Heading 4')
doc.add_paragraph(
    '유효한 세그먼트에 대해 4차 Butterworth 대역통과 필터(zero-phase, filtfilt)를 적용하였다:'
)
filter_info = [
    'PPG (PLETH): 0.5 ~ 8.0 Hz — 심장 박동 주파수 대역을 보존하면서 기저선 변동과 고주파 노이즈를 제거',
    'ABP (IBP1): 0.5 ~ 40.0 Hz — 혈압 파형의 고조파 성분을 보존하기 위해 더 넓은 대역 사용',
]
for f in filter_info:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph('3.2.4 결측값 처리', style='Heading 4')
doc.add_paragraph(
    '세그먼트 내 소수의 NaN 값(전체의 5% 미만)은 선형 보간(linear interpolation)으로 '
    '대체하였다.'
)

doc.add_paragraph('3.2.5 혈압 레이블', style='Heading 4')
doc.add_paragraph(
    '각 10초 세그먼트의 목표값(ground truth)은 해당 구간의 ART1_SBP, ART1_DBP, ART1_MBP '
    '중앙값(median)을 사용하였다. 추가적인 생리학적 필터링으로 SBP < 30 또는 > 250 mmHg, '
    'DBP < 10 또는 > 200 mmHg, SBP ≤ DBP인 세그먼트는 제외하였다.'
)

add_heading_styled('3.3 특징 추출 + 머신러닝 (Step 2)', level=2)

doc.add_paragraph('3.3.1 PPG 특징 추출', style='Heading 4')
doc.add_paragraph(
    '각 10초 PPG 세그먼트에서 총 37개의 특징을 추출하였다. 추출된 특징은 다음과 같이 '
    '분류된다:'
)

feat_headers = ['카테고리', '특징명', '설명']
feat_rows = [
    ['기본 통계 (6개)', 'mean, std, skewness,\nkurtosis, ptp, rms', '평균, 표준편차, 왜도, 첨도,\npeak-to-peak, RMS'],
    ['피크 관련 (8개)', 'n_peaks, heart_rate,\nppi_mean/std/cv,\npeak_amp_mean/std', '피크 수, 심박수,\n피크간 간격 통계,\n피크 진폭 통계'],
    ['골 관련 (2개)', 'valley_amp_mean/std', '골(valley) 진폭 통계'],
    ['맥파 진폭 (2개)', 'pulse_amp_mean/std', '피크-골 진폭차 통계'],
    ['시간 특징 (3개)', 'sys_time_mean,\ndia_time_mean,\nsys_dia_ratio', '수축기/이완기 시간,\n비율'],
    ['면적 (1개)', 'auc_total', '파형 하 면적'],
    ['미분 특징 (6개)', 'vpg_max/min/std,\napg_max/min/std', '1차 미분(VPG),\n2차 미분(APG) 통계'],
    ['주파수 특징 (7개)', 'vlf/lf/hf/cardiac_power,\nlf_hf_ratio, cardiac_ratio,\ndominant_freq,\nspectral_entropy', 'VLF/LF/HF/심장 대역 파워,\nLF/HF 비율, 심장 대역 비율,\n지배 주파수, 스펙트럼 엔트로피'],
    ['인구통계 (2개)', 'age, sex', '나이, 성별 (M=1, F=0)'],
]
add_table_with_style(feat_headers, feat_rows)
doc.add_paragraph()

doc.add_paragraph('3.3.2 머신러닝 모델', style='Heading 4')
doc.add_paragraph('두 가지 Gradient Boosting 모델을 사용하였다:')

ml_headers = ['하이퍼파라미터', 'XGBoost', 'LightGBM']
ml_rows = [
    ['n_estimators', '300', '300'],
    ['max_depth', '6', '6'],
    ['learning_rate', '0.05', '0.05'],
    ['subsample', '0.8', '0.8'],
    ['colsample_bytree', '0.8', '0.8'],
    ['reg_alpha', '0.1', '0.1'],
    ['reg_lambda', '1.0', '1.0'],
]
add_table_with_style(ml_headers, ml_rows)
doc.add_paragraph()
doc.add_paragraph(
    '특징 스케일링은 StandardScaler (z-score 정규화)를 적용하였으며, 각 fold마다 '
    '학습 데이터에 대해 fit 후 테스트 데이터에 transform하여 데이터 누출을 방지하였다.'
)

add_heading_styled('3.4 1D-CNN 딥러닝 모델 (Step 3)', level=2)

doc.add_paragraph('3.4.1 모델 구조', style='Heading 4')
doc.add_paragraph(
    '원시 PPG 파형을 직접 입력으로 사용하는 1D Convolutional Neural Network를 설계하였다. '
    '모델은 크게 세 부분으로 구성된다: (1) 합성곱 특징 추출기, (2) 인구통계 정보 결합, '
    '(3) 완전연결 예측 헤드.'
)

cnn_headers = ['Layer', 'Output Shape', 'Parameters']
cnn_rows = [
    ['Input', '(1, 1250)', '-'],
    ['Conv1d(1→16, k=15) + BN + ReLU + MaxPool(4)', '(16, 312)', '256'],
    ['Conv1d(16→32, k=7) + BN + ReLU + MaxPool(4)', '(32, 78)', '3,616'],
    ['Conv1d(32→64, k=5) + BN + ReLU + MaxPool(4)', '(64, 19)', '10,368'],
    ['Conv1d(64→128, k=3) + BN + ReLU', '(128, 19)', '24,832'],
    ['AdaptiveAvgPool1d(1)', '(128, 1)', '-'],
    ['Flatten + Concat [age, sex]', '(130,)', '-'],
    ['Linear(130→64) + ReLU + Dropout(0.3)', '(64,)', '8,384'],
    ['Linear(64→3)', '(3,)', '195'],
    ['Output: [SBP, DBP, MBP]', '(3,)', '-'],
]
add_table_with_style(cnn_headers, cnn_rows)
doc.add_paragraph()

doc.add_paragraph('3.4.2 Late Fusion', style='Heading 4')
doc.add_paragraph(
    '인구통계 정보(나이, 성별)는 Late Fusion 방식으로 결합하였다. 즉, CNN의 합성곱 층을 '
    '통과한 후 Global Average Pooling으로 얻은 특징 벡터에 나이와 성별을 연결(concatenate)하여 '
    '완전연결층의 입력으로 사용하였다. 나이와 성별은 StandardScaler로 정규화하였다.'
)

doc.add_paragraph('3.4.3 학습 설정', style='Heading 4')
train_headers = ['항목', '설정값']
train_rows = [
    ['손실 함수', 'Smooth L1 Loss (Huber Loss)'],
    ['옵티마이저', 'Adam (weight_decay=1e-4)'],
    ['학습률', '2e-3'],
    ['학습률 스케줄러', 'ReduceLROnPlateau (patience=2, factor=0.5)'],
    ['배치 크기', '256'],
    ['최대 에폭', '20'],
    ['조기 종료', 'Patience=5 (검증 손실 기준)'],
    ['Gradient Clipping', 'max_norm=1.0'],
    ['PPG 정규화', '학습 데이터 평균/표준편차로 z-score 정규화'],
]
add_table_with_style(train_headers, train_rows)
doc.add_paragraph()

add_heading_styled('3.5 교차 검증 전략', level=2)
doc.add_paragraph(
    '혈압 예측 모델의 일반화 성능을 정확하게 평가하기 위해 환자 단위(subject-level) 교차 검증을 '
    '적용하였다. 이는 같은 환자의 세그먼트가 학습 세트와 테스트 세트에 동시에 포함되는 것을 '
    '방지하여 데이터 누출(data leakage)을 원천 차단한다.'
)

cv_headers = ['방법', '적용 모델', '설명']
cv_rows = [
    ['Leave-One-Subject-Out\n(LOSO)', 'XGBoost,\nLightGBM', '46명 중 1명을 테스트로 사용,\n나머지 45명으로 학습.\n총 46 fold 수행.'],
    ['5-Fold Group K-Fold', '1D-CNN', '46명을 5개 그룹으로 분할,\n각 fold에서 ~9명 테스트.\nCPU 학습 효율을 위해 선택.\n환자당 최대 500개 세그먼트 서브샘플링.'],
]
add_table_with_style(cv_headers, cv_rows)
doc.add_paragraph()

add_heading_styled('3.6 평가 지표', level=2)
doc.add_paragraph('다음 평가 지표를 사용하여 모델 성능을 정량적으로 평가하였다:')

metric_headers = ['지표', '수식 / 기준', '의미']
metric_rows = [
    ['MAE', '|pred - true| 의 평균', '평균 절대 오차 (mmHg)'],
    ['RMSE', '√(MSE)', '제곱근 평균 제곱 오차'],
    ['R²', '1 - SS_res/SS_tot', '결정계수 (설명력)'],
    ['Bias ± SD', 'mean(diff) ± std(diff)', 'Bland-Altman 분석'],
    ['BHS Grade', '≤5/≤10/≤15 mmHg 비율', 'British Hypertension Society\n기준 등급'],
    ['AAMI', 'MAE≤5, SD≤8', 'Association for the Advancement\nof Medical Instrumentation'],
]
add_table_with_style(metric_headers, metric_rows)
doc.add_paragraph()

doc.add_paragraph('BHS (British Hypertension Society) 등급 기준:')
bhs_headers = ['Grade', '≤5 mmHg', '≤10 mmHg', '≤15 mmHg']
bhs_rows = [
    ['A', '≥60%', '≥85%', '≥95%'],
    ['B', '≥50%', '≥75%', '≥90%'],
    ['C', '≥40%', '≥65%', '≥85%'],
    ['D', '위 기준 미달', '-', '-'],
]
add_table_with_style(bhs_headers, bhs_rows)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 4. RESULTS
# ════════════════════════════════════════════════════════════════
add_heading_styled('4. 결과 (Results)', level=1)

add_heading_styled('4.1 데이터 추출 결과', level=2)
doc.add_paragraph(
    '46명의 환자에서 총 83,439개의 유효한 10초 세그먼트가 추출되었다. 환자당 평균 1,814개 '
    '(범위: 147~4,239개)의 세그먼트가 생성되었으며, 이는 녹화 시간과 신호 품질에 따라 '
    '편차가 있다.'
)

extract_headers = ['항목', '값']
extract_rows = [
    ['총 세그먼트 수', '83,439'],
    ['유효 세그먼트 수 (필터링 후)', '81,812'],
    ['환자당 평균 세그먼트', '1,814 (범위: 147~4,239)'],
    ['세그먼트 길이', '10초 (1,250 샘플 @ 125Hz)'],
    ['SBP 범위', '30 ~ 232 mmHg (평균 105 ± 18)'],
    ['DBP 범위', '10 ~ 194 mmHg (평균 55 ± 11)'],
]
add_table_with_style(extract_headers, extract_rows)
doc.add_paragraph()

# Per-case stats (condensed)
add_heading_styled('표 2. 환자별 데이터 추출 요약', level=3)
case_stats = [
    ['case01','F','1','1,910','1,892','75','43'],
    ['case02','M','12','2,855','2,848','114','55'],
    ['case03','M','12','1,982','1,957','127','64'],
    ['case04','F','10','830','821','94','55'],
    ['case05','M','8','2,042','2,000','93','50'],
    ['case06','M','3','1,421','1,402','89','49'],
    ['case07','M','15','1,523','1,493','106','45'],
    ['case08','M','2','2,048','2,030','96','46'],
    ['case09','M','15','2,778','2,750','118','49'],
    ['case10','M','4','2,002','1,996','107','70'],
    ['case11','M','5','1,416','1,414','101','49'],
    ['case12','M','5','3,150','3,091','110','57'],
    ['case13','F','0.25','1,010','949','59','37'],
    ['case14','F','7','1,118','1,033','105','49'],
    ['case15','M','2','1,297','1,247','87','56'],
    ['case16','F','3','3,296','3,210','111','49'],
    ['case17','M','9','971','963','102','65'],
    ['case18','M','9','3,227','3,196','100','63'],
    ['case19','M','3','936','923','92','47'],
    ['case20','M','12','4,184','4,157','118','56'],
    ['case21','M','3','4,239','4,129','109','46'],
    ['case22','M','5','865','865','102','58'],
    ['case23','F','7','3,375','3,339','109','55'],
    ['case24','M','9','877','874','89','44'],
    ['case25','M','5','2,578','2,547','106','58'],
    ['case26','M','2','1,136','1,126','101','54'],
    ['case27','M','9','2,597','2,564','120','60'],
    ['case28','M','1','3,786','3,753','105','52'],
    ['case29','M','2','2,921','2,698','119','63'],
    ['case30','F','7','1,227','1,221','119','60'],
    ['case31','M','5','147','143','111','49'],
    ['case32','M','7','505','495','91','57'],
    ['case33','M','7','1,591','1,574','117','65'],
    ['case34','F','7','1,249','1,175','113','69'],
    ['case35','M','8','1,237','1,194','106','61'],
    ['case36','F','4','1,214','1,107','99','57'],
    ['case37','F','7','1,222','1,206','105','65'],
    ['case38','F','4','1,154','1,117','95','70'],
    ['case39','M','12','866','859','122','68'],
    ['case40','F','10','1,080','1,071','119','56'],
    ['case41','F','10','1,389','1,346','89','55'],
    ['case42','M','5','1,221','1,175','115','68'],
    ['case43','F','7','984','974','89','56'],
    ['case44','M','7','700','676','82','50'],
    ['case45','F','13','1,083','1,077','87','48'],
    ['case46','F','5','4,200','4,135','95','53'],
]
add_table_with_style(
    ['Case', 'Sex', 'Age', 'Segments', 'Valid', 'SBP\n(mean)', 'DBP\n(mean)'],
    case_stats
)
doc.add_paragraph()

doc.add_page_break()

add_heading_styled('4.2 머신러닝 모델 결과', level=2)
doc.add_paragraph(
    '37개 PPG 특징 + 인구통계 정보(나이, 성별)를 입력으로 사용하여 XGBoost와 LightGBM 모델을 '
    'Leave-One-Subject-Out (LOSO) 교차 검증으로 평가하였다.'
)

add_heading_styled('표 3. 머신러닝 모델 성능 (LOSO CV)', level=3)
ml_result_headers = ['Target', 'Model', 'MAE\n(mmHg)', 'RMSE\n(mmHg)', 'R²', 'Bias±SD\n(mmHg)', 'BHS\nGrade']
ml_result_rows = [
    ['SBP', 'XGBoost', '14.98', '19.47', '-0.1925', '+0.88±19.45', 'D'],
    ['SBP', 'LightGBM', '16.09', '20.62', '-0.3375', '-0.23±20.62', 'D'],
    ['DBP', 'XGBoost', '8.66', '11.62', '-0.0780', '+1.76±11.49', 'D'],
    ['DBP', 'LightGBM', '8.42', '11.33', '-0.0249', '+1.56±11.22', 'D'],
    ['MBP', 'XGBoost', '10.41', '13.96', '-0.1436', '+1.72±13.85', 'D'],
    ['MBP', 'LightGBM', '10.15', '13.65', '-0.0935', '+1.43±13.58', 'D'],
]
add_table_with_style(ml_result_headers, ml_result_rows)
doc.add_paragraph()

doc.add_paragraph(
    '머신러닝 모델의 R²가 모두 음수인 것은 환자 간 혈압 변이가 크고, PPG 특징만으로는 '
    '개별 환자의 절대 혈압 수준을 예측하기 어렵다는 것을 의미한다. 그러나 Bias가 작은 것은 '
    '전체적인 평균 예측이 편향되지 않음을 보여준다.'
)

add_heading_styled('4.3 1D-CNN 모델 결과', level=2)
doc.add_paragraph(
    '원시 PPG 파형을 직접 입력으로 사용하는 1D-CNN 모델을 5-Fold Group K-Fold 교차 검증으로 '
    '평가하였다. 학습 효율을 위해 환자당 최대 500개 세그먼트를 랜덤 서브샘플링하여 총 22,638개 '
    '샘플을 사용하였다.'
)

add_heading_styled('표 4. 1D-CNN 모델 성능 (5-Fold Group CV)', level=3)
cnn_result_headers = ['Target', 'MAE\n(mmHg)', 'RMSE\n(mmHg)', 'R²', 'Bias±SD\n(mmHg)', '≤5\nmmHg', '≤10\nmmHg', '≤15\nmmHg', 'BHS']
cnn_result_rows = [
    ['SBP', '12.95', '16.76', '0.1569', '+0.5±16.8', '25.8%', '47.8%', '65.4%', 'D'],
    ['DBP', '7.47', '10.21', '0.2046', '-0.1±10.2', '42.8%', '74.8%', '90.6%', 'C'],
    ['MBP', '8.49', '11.67', '0.2193', '-0.1±11.7', '39.1%', '68.8%', '84.4%', 'D'],
]
add_table_with_style(cnn_result_headers, cnn_result_rows)
doc.add_paragraph()

add_heading_styled('표 5. 1D-CNN Fold별 성능', level=3)
fold_headers = ['Fold', 'Test\nSubjects', 'Test\nSamples', 'SBP MAE', 'DBP MAE', 'MBP MAE', 'Epochs']
fold_rows = [
    ['1', '9', '4,500', '15.06', '7.56', '9.24', '12'],
    ['2', '9', '4,500', '13.24', '6.94', '8.41', '16'],
    ['3', '9', '4,500', '10.37', '7.69', '7.52', '14'],
    ['4', '9', '4,500', '13.00', '7.69', '8.69', '12'],
    ['5', '10', '4,638', '13.09', '7.49', '8.60', '10'],
]
add_table_with_style(fold_headers, fold_rows)
doc.add_paragraph()

add_heading_styled('4.4 모델 비교', level=2)
doc.add_paragraph(
    '1D-CNN이 모든 예측 대상(SBP, DBP, MBP)에서 XGBoost 대비 개선된 성능을 보였다. '
    '특히 R²가 음수에서 양수로 전환되어 모델이 환자 간 변이를 일정 부분 설명할 수 있게 되었다.'
)

add_heading_styled('표 6. 모델 간 비교 요약', level=3)
comp_headers = ['Target', 'Model', 'CV 방법', 'MAE\n(mmHg)', 'RMSE\n(mmHg)', 'R²']
comp_rows = [
    ['SBP', 'XGBoost', 'LOSO (46-fold)', '14.98', '19.47', '-0.1925'],
    ['SBP', '1D-CNN', 'Group 5-fold', '12.95', '16.76', '+0.1569'],
    ['DBP', 'XGBoost', 'LOSO (46-fold)', '8.66', '11.62', '-0.0780'],
    ['DBP', '1D-CNN', 'Group 5-fold', '7.47', '10.21', '+0.2046'],
    ['MBP', 'XGBoost', 'LOSO (46-fold)', '10.41', '13.96', '-0.1436'],
    ['MBP', '1D-CNN', 'Group 5-fold', '8.49', '11.67', '+0.2193'],
]
add_table_with_style(comp_headers, comp_rows)
doc.add_paragraph()

doc.add_paragraph(
    '주요 개선 사항: SBP MAE 14.98→12.95 mmHg (13.5% 감소), '
    'DBP MAE 8.66→7.47 mmHg (13.7% 감소), '
    'MBP MAE 10.41→8.49 mmHg (18.4% 감소).'
)

add_heading_styled('4.5 특징 중요도 분석', level=2)
doc.add_paragraph(
    'XGBoost 모델의 특징 중요도(feature importance) 분석 결과, 인구통계 정보인 '
    '성별(sex)과 나이(age)가 가장 높은 중요도를 보였다. 이는 소아 환자에서 연령과 성별에 따른 '
    '혈압 차이가 매우 크다는 것을 반영한다.'
)

fi_headers = ['순위', '특징명', '중요도', '카테고리']
fi_rows = [
    ['1', 'sex (성별)', '0.2166', '인구통계'],
    ['2', 'age (나이)', '0.1177', '인구통계'],
    ['3', 'dominant_freq', '0.0546', '주파수'],
    ['4', 'peak_amp_mean', '0.0544', '피크'],
    ['5', 'apg_min', '0.0445', '미분'],
    ['6', 'skewness', '0.0402', '통계'],
    ['7', 'spectral_entropy', '0.0395', '주파수'],
    ['8', 'ppi_cv', '0.0360', '피크'],
    ['9', 'sys_time_mean', '0.0335', '시간'],
    ['10', 'apg_std', '0.0323', '미분'],
    ['11', 'ppi_mean', '0.0284', '피크'],
    ['12', 'cardiac_ratio', '0.0261', '주파수'],
    ['13', 'dia_time_mean', '0.0226', '시간'],
    ['14', 'vpg_max', '0.0226', '미분'],
    ['15', 'sys_dia_ratio', '0.0218', '시간'],
]
add_table_with_style(fi_headers, fi_rows)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 5. DISCUSSION
# ════════════════════════════════════════════════════════════════
add_heading_styled('5. 고찰 (Discussion)', level=1)

add_heading_styled('5.1 결과 해석', level=2)
doc.add_paragraph(
    '본 연구에서 1D-CNN 모델은 DBP 예측에서 BHS Grade C를 달성하였으며, 이는 '
    'PPG 파형으로부터 이완기혈압을 어느 정도 예측할 수 있음을 시사한다. 반면 SBP 예측은 '
    'MAE 12.95 mmHg로 임상적 허용 기준(AAMI: MAE ≤ 5 mmHg)에는 미치지 못하였다.'
)
doc.add_paragraph(
    '1D-CNN이 전통적 특징 추출 + XGBoost 대비 일관되게 우수한 성능을 보인 것은, '
    '수작업 특징으로 포착하지 못하는 PPG 파형의 미세한 패턴을 CNN이 자동으로 학습할 수 '
    '있음을 보여준다. 특히 R²가 음수에서 양수로 전환된 것은 딥러닝이 환자 간 변이를 '
    '일정 부분 모델링할 수 있음을 의미한다.'
)
doc.add_paragraph(
    '특징 중요도 분석에서 성별과 나이가 가장 높은 중요도를 보인 것은 소아 환자의 '
    '혈역학적 특성이 연령과 성별에 크게 의존한다는 기존 의학적 지식과 일치한다. '
    '특히 연구 대상의 연령 범위가 0.25~15세로 넓어 혈압의 정상 범위 자체가 '
    '연령에 따라 크게 다르다.'
)

add_heading_styled('5.2 한계점', level=2)
limitations = [
    '데이터 규모: 46명의 소아 환자로 구성된 비교적 작은 데이터셋을 사용하였다. '
    '더 많은 환자 데이터를 확보하면 모델의 일반화 성능이 향상될 수 있다.',

    '개인 보정 부재: 현재 모델은 환자별 보정(calibration) 없이 절대 혈압값을 '
    '예측한다. 실제 임상에서는 초기 NIBP 측정값으로 개인 보정을 수행하면 '
    '상당한 성능 향상이 기대된다.',

    '수술 중 데이터: 수술 중에는 마취, 약물, 체위 변경 등의 영향으로 PPG와 ABP의 '
    '관계가 정상 상태와 다를 수 있다. 다양한 임상 환경에서의 검증이 필요하다.',

    'GPU 미활용: CPU에서의 학습으로 인해 모델 크기와 데이터 사용량에 제한이 있었다. '
    'GPU 학습 환경이 구축되었으므로 향후 더 큰 모델과 전체 데이터를 활용할 수 있다.',

    '파형 재구성 미수행: 현재는 SBP/DBP/MBP 수치 예측만 수행하였으며, '
    'ABP 파형 전체를 재구성하는 접근법은 향후 과제이다.',
]
for lim in limitations:
    doc.add_paragraph(lim, style='List Number')

add_heading_styled('5.3 향후 연구 방향', level=2)
future = [
    '개인 보정(Personalized Calibration): 환자별 초기 혈압 측정값을 활용하여 '
    '모델의 offset을 보정하는 방법 개발. 이는 가장 큰 성능 향상이 기대되는 접근법이다.',

    'U-Net 기반 파형 재구성: PPG 파형으로부터 ABP 파형 전체를 재구성하는 1D U-Net 모델 개발. '
    '파형 재구성 후 SBP/DBP를 추출하면 보다 풍부한 혈역학적 정보를 얻을 수 있다.',

    'ResNet1D / Transformer: 더 깊고 강력한 딥러닝 아키텍처 탐색. Residual connection과 '
    'self-attention 메커니즘이 PPG 파형의 장거리 의존성 포착에 도움이 될 수 있다.',

    '다중 생체 신호 활용: ECG, SpO2 등 추가 생체 신호를 보조 입력으로 활용하여 '
    '예측 정확도 향상.',

    '앙상블 학습: ML 모델과 CNN 모델의 예측을 결합하여 상호 보완적인 성능 개선.',

    '외부 검증: 다른 의료기관의 소아 환자 데이터로 외부 검증(external validation)을 '
    '수행하여 일반화 가능성 확인.',
]
for f in future:
    doc.add_paragraph(f, style='List Number')


doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ════════════════════════════════════════════════════════════════
add_heading_styled('6. 결론 (Conclusion)', level=1)
doc.add_paragraph(
    '본 연구에서는 46명의 소아 환자로부터 수집된 수술 중 PPG 파형을 이용하여 '
    '침습적 동맥혈압(SBP, DBP, MBP)을 예측하는 모델을 개발하고 비교 분석하였다.'
)
doc.add_paragraph(
    '주요 결과는 다음과 같다:'
)
conclusions = [
    '총 83,439개의 유효한 10초 PPG-ABP 세그먼트 쌍을 추출하였다.',
    '37개의 PPG 특징을 활용한 XGBoost 모델은 LOSO 교차 검증에서 '
    'SBP MAE 14.98 mmHg, DBP MAE 8.66 mmHg를 기록하였다.',
    '원시 PPG 파형을 입력으로 하는 1D-CNN 모델은 5-Fold Group CV에서 '
    'SBP MAE 12.95 mmHg, DBP MAE 7.47 mmHg를 달성하여 머신러닝 모델 대비 '
    '13~18%의 성능 향상을 보였다.',
    'DBP 예측에서 BHS Grade C (≤5mmHg: 42.8%, ≤10mmHg: 74.8%, ≤15mmHg: 90.6%)를 '
    '달성하였다.',
    '특징 중요도 분석에서 성별과 나이가 가장 중요한 예측 변수로 확인되어, '
    '소아 환자의 혈압 예측에서 인구통계 정보의 필수적 역할을 입증하였다.',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Number')

doc.add_paragraph(
    '비록 현재 모델의 성능이 임상 표준(AAMI: MAE ≤ 5 mmHg, SD ≤ 8 mmHg)에는 미치지 '
    '못하지만, 소아 환자 대상의 PPG 기반 혈압 예측 연구의 기초를 마련하였으며, '
    '개인 보정, 더 강력한 모델 아키텍처, 더 많은 데이터 확보를 통해 성능 향상의 '
    '여지가 충분하다.'
)


# ════════════════════════════════════════════════════════════════
# 7. REFERENCES
# ════════════════════════════════════════════════════════════════
add_heading_styled('7. 참고문헌 (References)', level=1)
refs = [
    'Slapničar, G., Mlakar, N., & Luštrek, M. (2019). Blood pressure estimation from '
    'photoplethysmogram using a spectro-temporal deep neural network. Sensors, 19(15), 3420.',

    'El-Hajj, C., & Kyriacou, P. A. (2020). A review of machine learning techniques in '
    'photoplethysmography for the non-invasive cuff-less measurement of blood pressure. '
    'Biomedical Signal Processing and Control, 58, 101870.',

    'Ibtehaz, N., et al. (2022). PPG2ABP: Translating photoplethysmogram (PPG) signals '
    'to arterial blood pressure (ABP) waveforms. Biomedical Signal Processing and Control, 73, 103413.',

    'Lee, H. C., & Jung, C. W. (2018). Vital Recorder—a free research tool for automatic '
    'recording of high-resolution time-synchronised physiological data from multiple anaesthesia '
    'devices. Scientific Reports, 8(1), 1527.',

    'O\'Brien, E., et al. (2010). European Society of Hypertension International Protocol '
    'revision 2010 for the validation of blood pressure measuring devices in adults. '
    'Blood Pressure Monitoring, 15(1), 23-38.',

    'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. '
    'Proceedings of the 22nd ACM SIGKDD, 785-794.',

    'National High Blood Pressure Education Program Working Group on High Blood Pressure '
    'in Children and Adolescents. (2004). The fourth report on the diagnosis, evaluation, '
    'and treatment of high blood pressure in children and adolescents. Pediatrics, 114(2), 555-576.',
]
for i, ref in enumerate(refs):
    doc.add_paragraph(f'[{i+1}] {ref}')


# ── Save ──
output_path = 'C:/Users/jaege/Desktop/Study/PPG2ABP/PPG2ABP_Report.docx'
doc.save(output_path)
print(f"Word report saved: {output_path}")
